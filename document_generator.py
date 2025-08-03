"""
Document Generator Module
Handles creation of Word and PDF documents from JSON data using templates
"""

import json
import os
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("Warning: python-docx not installed. Install with: pip install python-docx")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    print("Warning: reportlab not installed. Install with: pip install reportlab")

class DocumentGenerator:
    def __init__(self, word_template_path="word_template.docx"):
        self.word_template_path = word_template_path
        self.setup_fonts()
    
    def setup_fonts(self):
        """Setup fonts for PDF generation"""
        try:
            # Try to register a font that supports Cyrillic
            # This is optional - reportlab has basic support for common characters
            pass
        except Exception:
            pass
    
    def generate_word_document(self, data, output_path, template_path=None):
        """
        Generate Word document from data using template
        
        Args:
            data (dict): Form data as JSON
            output_path (str): Path to save the output document
            template_path (str): Path to Word template (optional)
        """
        try:
            template_path = template_path or self.word_template_path
            
            # Try to load existing template
            if os.path.exists(template_path):
                doc = Document(template_path)
                self._fill_word_template(doc, data)
            else:
                # Create new document if template doesn't exist
                doc = self._create_word_document_from_scratch(data)
            
            # Save document
            doc.save(output_path)
            
        except Exception as e:
            raise Exception(f"Error generating Word document: {str(e)}")
    
    def _fill_word_template(self, doc, data):
        """
        Fill Word template with data by replacing placeholders
        
        Placeholders format: {{field_name}}
        """
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_placeholders_in_paragraph(paragraph, data)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(paragraph, data)
    
    def _replace_placeholders_in_paragraph(self, paragraph, data):
        """Replace placeholders in a single paragraph"""
        full_text = paragraph.text
        
        # Find and replace all placeholders
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))
        
        # Update paragraph text if changes were made
        if full_text != paragraph.text:
            paragraph.clear()
            paragraph.add_run(full_text)
    
    def _create_word_document_from_scratch(self, data):
        """Create a new Word document from scratch"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('EEG Score Checklist Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()  # Empty line
        
        # Group data by sections (based on common field prefixes)
        sections = self._group_data_by_sections(data)
        
        for section_name, section_data in sections.items():
            # Add section heading
            doc.add_heading(section_name, level=1)
            
            # Add section data as table
            if section_data:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Field'
                hdr_cells[1].text = 'Value'
                
                # Data rows
                for key, value in section_data.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = self._format_field_name(key)
                    row_cells[1].text = str(value) if value is not None else ''
            
            doc.add_paragraph()  # Empty line between sections
        
        return doc
    
    def _group_data_by_sections(self, data):
        """Group form data by logical sections"""
        sections = {
            'Patient Information': {},
            'Background Activity': {},
            'Epileptiform Activity': {},
            'Provocation Responses': {},
            'Conclusion': {},
            'Other': {}
        }
        
        # Map field prefixes to sections
        field_mapping = {
            'patient_': 'Patient Information',
            'study_': 'Patient Information',
            'indication': 'Patient Information',
            'symmetry': 'Background Activity',
            'predominant_': 'Background Activity',
            'reactivity': 'Background Activity',
            'organization': 'Background Activity',
            'continuity': 'Background Activity',
            'epileptiform_': 'Epileptiform Activity',
            'hv_': 'Provocation Responses',
            'photo_': 'Provocation Responses',
            'impression': 'Conclusion'
        }
        
        for key, value in data.items():
            section_found = False
            for prefix, section in field_mapping.items():
                if key.startswith(prefix) or key == prefix:
                    sections[section][key] = value
                    section_found = True
                    break
            
            if not section_found:
                sections['Other'][key] = value
        
        # Remove empty sections
        return {k: v for k, v in sections.items() if v}
    
    def _format_field_name(self, field_name):
        """Format field name for display"""
        # Convert snake_case to Title Case
        formatted = field_name.replace('_', ' ').title()
        
        # Special formatting for common terms
        replacements = {
            'Hv': 'Hyperventilation',
            'Photo': 'Photostimulation',
            'Eeg': 'EEG'
        }
        
        for old, new in replacements.items():
            formatted = formatted.replace(old, new)
        
        return formatted
    
    def generate_pdf_document(self, data, output_path):
        """
        Generate PDF document from data
        
        Args:
            data (dict): Form data as JSON
            output_path (str): Path to save the PDF document
        """
        try:
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                alignment=1,  # Center alignment
                spaceAfter=30
            )
            story.append(Paragraph("EEG Score Checklist Report", title_style))
            
            # Generation date
            date_style = ParagraphStyle(
                'DateStyle',
                parent=styles['Normal'],
                fontSize=10,
                alignment=2,  # Right alignment
                spaceAfter=20
            )
            story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", date_style))
            
            # Group data by sections
            sections = self._group_data_by_sections(data)
            
            for section_name, section_data in sections.items():
                if not section_data:
                    continue
                
                # Section heading
                story.append(Paragraph(section_name, styles['Heading2']))
                story.append(Spacer(1, 12))
                
                # Create table data
                table_data = [['Field', 'Value']]  # Header row
                
                for key, value in section_data.items():
                    formatted_key = self._format_field_name(key)
                    formatted_value = str(value) if value is not None else ''
                    
                    # Handle long text values
                    if len(formatted_value) > 50:
                        formatted_value = formatted_value[:47] + "..."
                    
                    table_data.append([formatted_key, formatted_value])
                
                # Create and style table
                table = Table(table_data, colWidths=[2.5*inch, 4*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                story.append(table)
                story.append(Spacer(1, 20))
            
            # Build PDF
            doc.build(story)
            
        except Exception as e:
            raise Exception(f"Error generating PDF document: {str(e)}")
    
    def create_word_template(self, template_path="word_template.docx"):
        """
        Create a sample Word template with placeholders
        
        Args:
            template_path (str): Path where to save the template
        """
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('EEG Score Checklist Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Instructions
            doc.add_paragraph("This is a template document. Use {{field_name}} format for placeholders.")
            doc.add_paragraph()
            
            # Patient Information Section
            doc.add_heading('Patient Information', level=1)
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Table Grid'
            
            patient_fields = [
                ('Patient Name:', '{{patient_name}}'),
                ('Age:', '{{patient_age}}'),
                ('Sex:', '{{patient_sex}}'),
                ('Study Date:', '{{study_date}}'),
                ('Indication:', '{{indication}}')
            ]
            
            for i, (label, placeholder) in enumerate(patient_fields):
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = placeholder
            
            doc.add_paragraph()
            
            # Background Activity Section
            doc.add_heading('Background Activity', level=1)
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Table Grid'
            
            bg_fields = [
                ('Symmetry:', '{{symmetry}}'),
                ('Predominant Frequency:', '{{predominant_frequency}}'),
                ('Reactivity:', '{{reactivity}}'),
                ('Organization:', '{{organization}}'),
                ('Continuity:', '{{continuity}}')
            ]
            
            for i, (label, placeholder) in enumerate(bg_fields):
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = placeholder
            
            doc.add_paragraph()
            
            # Epileptiform Activity Section
            doc.add_heading('Epileptiform Activity', level=1)
            doc.add_paragraph(f"Present: {{{{epileptiform_present}}}}")
            doc.add_paragraph(f"Frequency: {{{{epileptiform_frequency}}}}")
            doc.add_paragraph(f"Location: {{{{epileptiform_location}}}}")
            doc.add_paragraph()
            
            # Provocation Responses Section
            doc.add_heading('Provocation Responses', level=1)
            doc.add_paragraph(f"Hyperventilation: {{{{hv_response}}}}")
            doc.add_paragraph(f"Photostimulation: {{{{photo_response}}}}")
            doc.add_paragraph()
            
            # Conclusion Section
            doc.add_heading('Conclusion', level=1)
            doc.add_paragraph("{{impression}}")
            
            # Save template
            doc.save(template_path)
            print(f"Word template created: {template_path}")
            
        except Exception as e:
            raise Exception(f"Error creating Word template: {str(e)}")

# Example usage and testing
if __name__ == "__main__":
    # Create sample data for testing
    sample_data = {
        "patient_name": "John Doe",
        "patient_age": 45,
        "patient_sex": "Male",
        "study_date": "2024-01-15",
        "indication": "Seizure evaluation",
        "symmetry": "Symmetric",
        "predominant_frequency": "Alpha",
        "reactivity": "Present",
        "organization": "Organized",
        "continuity": "Continuous",
        "epileptiform_present": True,
        "epileptiform_frequency": 2,
        "epileptiform_location": "Left temporal",
        "hv_response": "No change",
        "photo_response": "Normal photoresponse",
        "impression": "Abnormal EEG with left temporal epileptiform activity."
    }
    
    # Create document generator
    generator = DocumentGenerator()
    
    # Create Word template
    try:
        generator.create_word_template()
        print("Word template created successfully!")
    except Exception as e:
        print(f"Error creating Word template: {e}")
    
    # Test document generation
    try:
        generator.generate_word_document(sample_data, "test_output.docx")
        print("Word document generated successfully!")
    except Exception as e:
        print(f"Error generating Word document: {e}")
    
    try:
        generator.generate_pdf_document(sample_data, "test_output.pdf")
        print("PDF document generated successfully!")
    except Exception as e:
        print(f"Error generating PDF document: {e}")